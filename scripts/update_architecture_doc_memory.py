from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


DOCX_PATH = Path("/Users/rafalaffeier/projects/asistente/documentos/AI_Assistant_Architecture_Professiona.docx")


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.add_run(f"- {item}")


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    return table


def main():
    document = Document(DOCX_PATH)

    document.add_page_break()
    document.add_heading("79. Local-First File Memory Architecture", level=1)
    document.add_paragraph(
        "The assistant memory should be designed as a local-first, file-based system. "
        "Instead of treating the relational database as the source of truth for long-term "
        "user memory, the primary memory record should be stored as human-readable "
        "Markdown files. This keeps the memory portable, inspectable, easy for the AI "
        "to load selectively, and independent from any single database engine."
    )
    document.add_paragraph(
        "The desktop application should maintain a local memory folder, and the server "
        "should synchronize encrypted copies of those files so the user can access their "
        "assistant memory from desktop, mobile, and future web clients."
    )

    document.add_heading("Memory Folder Layout", level=2)
    add_table(
        document,
        ["Path", "Purpose"],
        [
            ("memory/journal/YYYY/MM/YYYY-MM-DD.md", "Daily journal with tasks, decisions, events, and pending items."),
            ("memory/summaries/YYYY-Www.md", "Weekly summaries generated from daily journals."),
            ("memory/summaries/YYYY-MM.md", "Monthly summaries for fast high-level recall."),
            ("memory/entities/projects.md", "Known projects, current status, owners, and context."),
            ("memory/entities/people.md", "User-approved people/entity memory."),
            ("memory/entities/preferences.md", "Stable user preferences and assistant behavior notes."),
            ("memory/index/*", "Rebuildable local search indexes and caches. Not the source of truth."),
        ],
    )

    document.add_heading("Daily Journal Contract", level=2)
    document.add_paragraph(
        "Each day should be represented by one Markdown file. The assistant writes concise, "
        "structured entries rather than unbounded raw logs."
    )
    add_bullets(
        document,
        [
            "Facts: important events that happened during the day.",
            "Tasks: work completed or started by the user or assistant.",
            "Decisions: architectural, product, technical, or personal decisions.",
            "Pending: follow-up items that should remain visible.",
            "Technical events: errors, deployments, environment changes, and relevant system state.",
        ],
    )

    document.add_heading("80. Server-Synchronized Markdown Memory", level=1)
    document.add_paragraph(
        "The server is still required for a functional multi-device product. However, its "
        "role should be to synchronize and authorize access to encrypted memory files, "
        "not to become the main long-term memory database."
    )
    add_table(
        document,
        ["Component", "Responsibility"],
        [
            ("Desktop App", "Maintains the full local memory folder and decrypts memory for local AI use."),
            ("Mobile App", "Reads synchronized memory, edits notes, and sends remote actions through the backend."),
            ("Backend API", "Handles auth, devices, permissions, sync metadata, and relay between clients."),
            ("Object Storage", "Stores encrypted Markdown files and attachments."),
            ("PostgreSQL", "Stores administrative metadata only: users, devices, sessions, versions, permissions, and sync state."),
        ],
    )

    document.add_heading("Sync Rules", level=2)
    add_bullets(
        document,
        [
            "Markdown files are the source of truth for assistant memory.",
            "Search indexes, embeddings, and local SQLite caches are rebuildable.",
            "Each file has a version, checksum, modified timestamp, and owning workspace.",
            "Clients upload encrypted file changes and download encrypted updates.",
            "Conflicts should be resolved per file, preserving both versions when automatic merge is unsafe.",
        ],
    )

    document.add_heading("81. Minimal Database Role", level=1)
    document.add_paragraph(
        "A relational database remains useful, but it should be intentionally small. "
        "It should support the product infrastructure rather than store the user's "
        "private memory content."
    )
    add_table(
        document,
        ["Database Area", "Recommended Use"],
        [
            ("users", "Account identity, login provider, and account status."),
            ("devices", "Registered desktop/mobile clients, trust level, revocation state."),
            ("sessions", "Short-lived access sessions and refresh token rotation metadata."),
            ("file_versions", "Path, owner, version, checksum, size, and encrypted storage reference."),
            ("sync_events", "Append-only metadata about uploads, downloads, merges, and conflicts."),
            ("permissions", "Workspace and device-level access rules."),
            ("billing", "Optional future subscription and plan metadata."),
        ],
    )
    document.add_paragraph(
        "The database must not be treated as the assistant's narrative memory. If an index "
        "or embedding store is added later, it should be considered derived data that can "
        "be deleted and rebuilt from the Markdown source files."
    )

    document.add_heading("82. End-to-End Encryption for Memory Files", level=1)
    document.add_paragraph(
        "Private memory files should be encrypted before upload. The server and object "
        "storage should persist encrypted bytes and minimal metadata. Plaintext should "
        "exist only on authorized user devices during active use."
    )
    add_table(
        document,
        ["Security Control", "Requirement"],
        [
            ("Client-side encryption", "Encrypt Markdown files before upload; decrypt only on authorized devices."),
            ("Per-user keys", "Use keys scoped to a user or workspace, never a single global application key."),
            ("Device key storage", "Store device secrets in macOS Keychain, Windows DPAPI/Credential Manager, iOS Keychain, or Android Keystore."),
            ("Signed URLs", "Use short-lived signed URLs for object storage access; never public memory URLs."),
            ("Transport security", "Require HTTPS for all API and storage operations."),
            ("Metadata minimization", "Avoid putting sensitive content in filenames, tags, logs, or database fields."),
            ("Audit events", "Log access metadata without storing decrypted memory content."),
        ],
    )

    document.add_heading("83. Recovery, Revocation, and Privacy Trade-Offs", level=1)
    document.add_paragraph(
        "End-to-end encryption creates an explicit recovery trade-off. If the service "
        "cannot read user memory, it also cannot recover that memory without a user-held "
        "recovery method. The product must make this clear and provide safe recovery flows."
    )
    add_bullets(
        document,
        [
            "Support device revocation so a lost laptop or phone can no longer sync memory.",
            "Provide a recovery kit or recovery phrase for restoring encrypted memory on a new device.",
            "Make it clear that losing all authorized devices and recovery material may make encrypted memory unrecoverable.",
            "Allow users to opt into encrypted cloud backup while preserving local Markdown export.",
            "Keep the product usable offline by writing locally first and syncing when connectivity returns.",
        ],
    )

    document.add_heading("Architecture Decision", level=2)
    document.add_paragraph(
        "The recommended architecture is: server required, database minimal, memory file-based, "
        "and private content encrypted end-to-end. This provides multi-device functionality "
        "without turning the user's long-term assistant memory into an opaque database."
    )

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
